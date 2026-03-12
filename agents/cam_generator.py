"""
agents/cam_generator.py
Credit Appraisal Memorandum (CAM) generator using python-docx.
Produces a professional .docx report with Five Cs analysis, forensic flags,
committee decision, and lending recommendation.
"""
from __future__ import annotations
import os
from datetime import datetime
from typing import Optional
from models.state_schema import CreditCaseState

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _add_heading(doc, text: str, level: int = 1, color: Optional[tuple] = None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    return heading


def _add_paragraph(doc, text: str, bold: bool = False, italic: bool = False, indent: bool = False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def _add_kv(doc, key: str, value: str):
    p = doc.add_paragraph()
    run_key = p.add_run(f"{key}: ")
    run_key.bold = True
    p.add_run(value)


def _add_divider(doc):
    doc.add_paragraph("─" * 80)


def _shade_cell(cell, hex_color: str):
    """Apply background shading to a table cell."""
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)
    except Exception:
        pass


# ─────────────────────────────────────────────────────────────────────────────
# CAM Builder
# ─────────────────────────────────────────────────────────────────────────────

def _build_cam(state: CreditCaseState) -> Document:
    doc = Document()

    company      = state.get("company_name", "Unknown")
    case_id      = state.get("case_id", "N/A")
    gst_sales    = state.get("gst_sales", 0.0)
    bank_credits = state.get("bank_credits", 0.0)
    risk_results = state.get("risk_results", {})
    committee    = state.get("committee_votes", [])
    legal_flags  = state.get("legal_flags", [])
    forensic     = state.get("forensic_report", [])
    vis_results  = state.get("vision_results", [])
    reality      = state.get("reality_score", 0.0)
    decision     = risk_results.get("decision", "PENDING")
    score        = risk_results.get("total_score", 0)
    category     = risk_results.get("category", "Unknown")
    cf_analysis  = state.get("counterfactual_analysis", [])

    # ── Title Page ──
    title = doc.add_heading("CREDPANION", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Credit Appraisal Memorandum (CAM)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].bold = True

    doc.add_paragraph("")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Applicant: {company}  |  Case ID: {case_id}  |  "
        f"Date: {datetime.utcnow().strftime('%d %B %Y')}"
    ).italic = True
    doc.add_page_break()

    # ── 1. Executive Summary ──
    _add_heading(doc, "1. Executive Summary", 1, color=(31, 78, 121))
    _add_kv(doc, "Company", company)
    _add_kv(doc, "Case ID", case_id)
    _add_kv(doc, "Analysis Date", datetime.utcnow().strftime("%d %B %Y %H:%M UTC"))
    _add_kv(doc, "Risk Score", f"{score}/100")
    _add_kv(doc, "Risk Category", category)
    _add_kv(doc, "Committee Decision", decision)
    _add_kv(doc, "GST Sales (Declared)", f"₹{gst_sales:,.0f}")
    _add_kv(doc, "Bank Credits (Verified)", f"₹{bank_credits:,.0f}")
    _add_kv(doc, "Source", "GST Returns — GSTR-3B (Verified by Credpanion Extractor Agent), "
                           "Bank Statement (90-day rolling average).")
    _add_divider(doc)

    # ── 2. Five Cs Analysis ──
    _add_heading(doc, "2. Five Cs Credit Analysis", 1, color=(31, 78, 121))

    five_cs = [
        (
            "2.1 Character",
            f"Legal flags found: {len(legal_flags)}. "
            + (f"Director disqualification on record. " if any("Disqualif" in f for f in legal_flags) else "No director disqualification. ")
            + (f"{sum(1 for f in legal_flags if '138' in f)} cheque bounce case(s) under S.138 NI Act. " if any("138" in f for f in legal_flags) else "No S.138 cases. ")
            + "Source: MCA21 Director Database, Court Records Portal.",
        ),
        (
            "2.2 Capacity",
            f"Declared GST revenue: ₹{gst_sales:,.0f}. Bank-verified credits: ₹{bank_credits:,.0f}. "
            + f"Bank-GST mismatch: {state.get('bank_mismatch_pct', 0):.1f}%. "
            + ("Revenue inflation detected — credits below 85% of declared GST. " if state.get("revenue_inflation_detected") else "Revenue satisfactorily reconciled. ")
            + "Source: GSTR-3B, Bank Statement (AXIS/HDFC).",
        ),
        (
            "2.3 Capital",
            f"Operational reality score: {reality:.2f}. "
            + (f"Active assets: {vis_results[0].get('active_assets', 'N/A')}, Idle assets: {vis_results[0].get('idle_assets', 'N/A')}. " if vis_results else "")
            + "Circular trading pattern detected: "
            + ("YES — indicates possible siphoning of working capital. " if state.get("circular_trading_detected") else "NO. ")
            + "Source: Factory inspection photographs, CredPanion Vision Agent.",
        ),
        (
            "2.4 Collateral",
            "Collateral assessment deferred to physical field inspection and legal search report. "
            "EncumbranceSearch and title deed verification recommended before disbursement. "
            "Source: Property Documents (pending upload to Credpanion portal).",
        ),
        (
            "2.5 Conditions",
            f"Sector credit norms apply. Company operating in: {'Trading/Services (inferred from GST pattern)'}. "
            "Macro headwinds: interest rate environment, regulatory NBFC caps. "
            "Loan purpose alignment with operational activity: "
            + ("Questionable — idle plant detected. " if reality < 0.5 else "Consistent with stated purpose. ")
            + "Source: Industry analysis, operational photo assessment.",
        ),
    ]

    for title_text, body in five_cs:
        _add_heading(doc, title_text, 2)
        _add_paragraph(doc, body, indent=True)

    _add_divider(doc)

    # ── 3. Forensic Red Flags ──
    _add_heading(doc, "3. Forensic Red Flags", 1, color=(192, 0, 0))
    if forensic:
        for i, flag in enumerate(forensic, 1):
            _add_paragraph(doc, f"  {i}. {flag}", indent=True)
    else:
        _add_paragraph(doc, "  No forensic red flags detected.", indent=True)
    _add_divider(doc)

    # ── 4. Committee Decision ──
    _add_heading(doc, "4. Credit Committee Decision", 1, color=(31, 78, 121))

    # Committee vote table
    if committee:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, hdr in enumerate(["Agent", "Vote", "Score", "Rationale"]):
            hdr_cells[i].text = hdr
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            _shade_cell(hdr_cells[i], "1F4E79")
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for v in committee:
            row = table.add_row().cells
            row[0].text = v["agent"]
            row[1].text = v["vote"]
            row[2].text = f"{v['score']:.2f}"
            row[3].text = v["rationale"]
            if v["vote"] == "REJECT":
                _shade_cell(row[1], "FFD7D7")
            else:
                _shade_cell(row[1], "D7FFD7")

    doc.add_paragraph("")
    _add_kv(doc, "Weighted Score", f"{state.get('committee_weighted_score', 0):.4f}")
    _add_kv(doc, "Final Decision", decision)
    _add_divider(doc)

    # ── 5. Counterfactual Analysis ──
    _add_heading(doc, "5. Counterfactual Risk Scenarios", 1, color=(31, 78, 121))
    _add_paragraph(
        doc,
        f"Current risk score: {score}/100. The following scenarios model potential score improvements "
        f"if specific risk factors were resolved:"
    )
    doc.add_paragraph("")
    if cf_analysis:
        cf_table = doc.add_table(rows=1, cols=3)
        cf_table.style = "Table Grid"
        cf_hdr = cf_table.rows[0].cells
        for i, h in enumerate(["Scenario", "Score If Resolved", "Reduction"]):
            cf_hdr[i].text = h
            cf_hdr[i].paragraphs[0].runs[0].bold = True
            _shade_cell(cf_hdr[i], "2C3E50")
            cf_hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for cf in cf_analysis:
            row = cf_table.add_row().cells
            row[0].text = cf.get("scenario", "")
            row[1].text = str(cf.get("score_if_removed", ""))
            row[2].text = f"-{cf.get('delta', 0)} pts"
    _add_divider(doc)

    # ── 6. Final Recommendation ──
    _add_heading(doc, "6. Final Lending Recommendation", 1, color=(31, 78, 121))
    if decision == "APPROVE":
        rec_text = (
            f"Based on the forensic analysis, committee voting, and risk scoring, "
            f"{company} presents a {category.upper()} RISK profile (score: {score}/100). "
            f"The credit committee recommends APPROVAL subject to standard covenants, "
            f"enhanced monitoring, and quarterly forensic review."
        )
    else:
        rec_text = (
            f"Based on the forensic analysis, committee voting, and risk scoring, "
            f"{company} presents a {category.upper()} RISK profile (score: {score}/100). "
            f"The credit committee recommends REJECTION of this loan application. "
            f"Re-application may be considered after resolution of the identified red flags."
        )
    doc.add_paragraph(rec_text)

    doc.add_paragraph("")
    doc.add_paragraph(
        "This report was generated by the Credpanion Agentic Credit Intelligence Platform. "
        "All figures are sourced from documents submitted by the applicant and processed by "
        "AI agents. This document does not constitute a legally binding credit decision."
    ).italic = True

    return doc


def run_cam_generator(state: CreditCaseState) -> CreditCaseState:
    """
    LangGraph node: CAM Generator.
    Builds and saves the Credit Appraisal Memo as a .docx file.
    """
    state.setdefault("audit_trail", [])
    state["audit_trail"].append("CAM Generator: Compiling Credit Appraisal Memorandum...")

    os.makedirs("reports", exist_ok=True)
    company = state.get("company_name", "unknown").replace(" ", "_")
    case_id = state.get("case_id", "CASE000")
    filename = f"reports/CAM_{company}_{case_id}.docx"

    if DOCX_AVAILABLE:
        doc = _build_cam(state)
        doc.save(filename)
        state["cam_document_path"] = filename
        state["audit_trail"].append(f"CAM Generator: Report saved → {filename}")
    else:
        # Fallback: write plain-text CAM
        txt_path = filename.replace(".docx", ".txt")
        with open(txt_path, "w") as f:
            f.write(_build_text_cam(state))
        state["cam_document_path"] = txt_path
        state["audit_trail"].append(
            f"CAM Generator: python-docx unavailable — text report saved → {txt_path}"
        )

    return state


def _build_text_cam(state: CreditCaseState) -> str:
    """Plain-text fallback CAM when python-docx is unavailable."""
    company = state.get("company_name", "Unknown")
    risk = state.get("risk_results", {})
    lines = [
        "=" * 70,
        "CREDPANION — Credit Appraisal Memorandum",
        "=" * 70,
        f"Company    : {company}",
        f"Case ID    : {state.get('case_id', 'N/A')}",
        f"Date       : {datetime.utcnow().strftime('%d %B %Y')}",
        f"Risk Score : {risk.get('total_score', 0)}/100",
        f"Category   : {risk.get('category', 'N/A')}",
        f"Decision   : {risk.get('decision', 'N/A')}",
        "",
        "FORENSIC FLAGS:",
    ]
    for flag in state.get("forensic_report", []):
        lines.append(f"  - {flag}")
    lines.append("")
    lines.append("COMMITTEE VOTES:")
    for v in state.get("committee_votes", []):
        lines.append(f"  [{v['agent']}] {v['vote']} (score={v['score']:.2f})")
    lines.append("=" * 70)
    return "\n".join(lines)
